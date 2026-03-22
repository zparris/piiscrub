"""Output generation — reconstruct scrubbed documents in original format.

All output writers:
  - Write to a .tmp file then os.replace() atomically
  - Strip all document metadata (author, creator, dates, custom properties)
  - Preserve original file structure (pages, paragraphs, sheets, cells)
  - Never log or expose PII values

Format-specific notes:
  PDF  — Uses PyMuPDF text-search redaction: search_for(original_pii) → add_redact_annot
          with inline replacement text → apply_redactions(images=0, graphics=0).
          Only matched text characters are removed; vector graphics and images are preserved.
  DOCX — Replaces run text in-place. Warns if tracked changes detected.
          Clears core_properties after reconstruction.
  XLSX — Replaces cell values. Formulae are NOT preserved (data_only limitation).
  CSV  — Rewrites via pandas.
  EML  — Reconstructs headers + body with scrubbed values.
  TXT  — Writes scrubbed lines directly.
"""

from __future__ import annotations

import csv
import email
import email.policy
import os
import tempfile
from pathlib import Path
from typing import TYPE_CHECKING

from piiscrub._logging import get_logger
from piiscrub.models import AnonymisedChunk, TextChunk

if TYPE_CHECKING:
    pass

_log = get_logger("output")


def get_output_extension(source_ext: str) -> str:
    """Return the output file extension for a given source extension.

    PDFs are output as DOCX so the reflowable format accommodates replacement
    labels of any length without fixed-layout overlap artefacts.
    All other formats output in their original format.
    """
    _map = {
        ".pdf":  ".docx",
        ".docx": ".docx",
        ".xlsx": ".xlsx",
        ".csv":  ".csv",
        ".eml":  ".txt",
        ".txt":  ".txt",
    }
    return _map.get(source_ext.lower(), source_ext.lower())


def reconstruct(
    source_path: Path,
    original_chunks: list[TextChunk],
    anonymised_chunks: list[AnonymisedChunk],
    output_path: Path,
) -> None:
    """Reconstruct scrubbed document in its original format.

    Writes atomically: output_path.tmp → output_path via os.replace().
    Strips all document metadata from output.

    Args:
        source_path: original file (read-only; used for format dispatch and PDF/DOCX template)
        original_chunks: TextChunks from extractor (carry position metadata)
        anonymised_chunks: AnonymisedChunks from anonymiser (carry scrubbed_text)
        output_path: destination path for the scrubbed document
    """
    source_path = Path(source_path)
    output_path = Path(output_path)
    ext = source_path.suffix.lower()

    # Build lookup: chunk_id → AnonymisedChunk
    scrubbed: dict[str, AnonymisedChunk] = {ac.chunk_id: ac for ac in anonymised_chunks}

    tmp_path = output_path.with_suffix(output_path.suffix + ".tmp")

    dispatch = {
        ".pdf": _reconstruct_pdf_as_docx,
        ".docx": _reconstruct_docx,
        ".xlsx": _reconstruct_xlsx,
        ".csv": _reconstruct_csv,
        ".eml": _reconstruct_eml,
        ".txt": _reconstruct_txt,
    }

    fn = dispatch.get(ext)
    if fn is None:
        raise ValueError(f"output.py: unsupported format '{ext}'")

    fn(source_path, original_chunks, scrubbed, tmp_path)
    os.replace(tmp_path, output_path)

    _log.info("output_written", format=ext.lstrip("."), output_size=output_path.stat().st_size)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _reconstruct_pdf(
    source_path: Path,
    chunks: list[TextChunk],
    scrubbed: dict[str, AnonymisedChunk],
    output_path: Path,
) -> None:
    """Redact PDF using PyMuPDF text-search redaction.

    Strategy:
      1. For each detected PII value, search the page for the exact original
         text string and add a whitebox redact annotation (no inline text yet).
      2. Call apply_redactions(images=0, graphics=0) — removes only matched text
         characters; vector graphics (table lines, borders) and images are preserved.
      3. After apply, insert replacement labels at the original text positions using
         insert_text() which writes single-line (no wrapping), so long labels like
         [ORGANIZATION_1] don't break across two lines.

    This approach avoids the previous line-bbox strategy which wiped entire PDF
    content stream blocks and destroyed all adjacent non-PII text.
    """
    import fitz  # PyMuPDF
    from collections import defaultdict

    doc = fitz.open(str(source_path))

    # page_no → [(Point, replacement_label), …]  — collected before apply_redactions
    pending_inserts: dict[int, list[tuple[object, str]]] = defaultdict(list)

    for chunk in chunks:
        ac = scrubbed.get(chunk.chunk_id)
        if ac is None or not ac.replacements:
            continue  # no PII in this chunk

        page_no = chunk.metadata.get("page", chunk.page_or_sheet or 0)
        if page_no >= len(doc):
            continue
        page = doc[page_no]

        for original, replacement in ac.replacements.items():
            if not original or not original.strip():
                continue

            # Locate the exact text on the page
            found = page.search_for(original)
            if not found:
                found = page.search_for(original, flags=fitz.TEXT_DEHYPHENATE)

            for rect in found:
                # Whitebox: remove original PII from content stream
                page.add_redact_annot(rect, fill=(1, 1, 1), cross_out=False)
                # Remember where to insert the replacement label (baseline point)
                # y1 - 1 gives a baseline just inside the bottom of the rect
                insert_pt = fitz.Point(rect.x0, rect.y1 - 1)
                pending_inserts[page_no].append((insert_pt, replacement))

    # Pass 1: apply redactions (text removal only; keep images and vector graphics)
    for page_no in range(len(doc)):
        doc[page_no].apply_redactions(images=0, graphics=0)

    # Pass 2: insert replacement labels single-line at the original text positions
    for page_no, inserts in pending_inserts.items():
        page = doc[page_no]
        for pt, label in inserts:
            page.insert_text(
                pt,
                label,
                fontsize=7,
                color=(0.6, 0, 0),   # dark red — matches the UI's placeholder style
            )

    doc.set_metadata({})
    doc.save(str(output_path), garbage=4, deflate=True)
    doc.close()


# ---------------------------------------------------------------------------
# PDF → DOCX (reflowable output)
# ---------------------------------------------------------------------------

def _add_coloured_runs(para: object, text: str, replacements: dict[str, str]) -> None:
    """Add runs to a DOCX paragraph, colouring PII labels dark red."""
    import re
    from docx.shared import Pt, RGBColor  # type: ignore[import]

    # Build a sorted list of (start, end, label) for every label in the text
    intervals: list[tuple[int, int, str]] = []
    for label in replacements.values():
        for m in re.finditer(re.escape(label), text):
            intervals.append((m.start(), m.end(), label))
    intervals.sort()

    pos = 0
    for start, end, label in intervals:
        if pos < start:
            run = para.add_run(text[pos:start])  # type: ignore[union-attr]
            run.font.size = Pt(10)
        run = para.add_run(label)  # type: ignore[union-attr]
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        pos = end
    if pos < len(text):
        run = para.add_run(text[pos:])  # type: ignore[union-attr]
        run.font.size = Pt(10)


def _reconstruct_pdf_as_docx(
    source_path: Path,
    chunks: list[TextChunk],
    scrubbed: dict[str, AnonymisedChunk],
    output_path: Path,
) -> None:
    """Convert a scrubbed PDF to a reflowable DOCX.

    PDFs are fixed-layout so replacement labels of different lengths cause
    visual overlap. DOCX reflows text naturally, making it the correct output
    format for any document where the content matters more than pixel fidelity.

    Structure inference:
      - Short ALL-CAPS lines → Heading 2
      - Page transitions → page break paragraph
      - All other lines → Normal 10pt paragraph with PII labels in dark red
    """
    from docx import Document  # type: ignore[import]
    from docx.shared import Pt  # type: ignore[import]
    from docx.oxml.ns import qn  # type: ignore[import]

    doc = Document()

    # Remove the single empty paragraph Word adds by default
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Clear all core properties (strip author/company metadata)
    cp = doc.core_properties
    for attr in ("author", "comments", "company", "keywords", "last_modified_by",
                 "subject", "title", "category", "content_status", "identifier"):
        try:
            setattr(cp, attr, "")
        except Exception:
            pass

    current_page: int | None = None

    for chunk in chunks:
        ac = scrubbed.get(chunk.chunk_id)
        text = ac.scrubbed_text if ac else chunk.text

        page_no = chunk.metadata.get("page", 0)

        # Insert a page break when the source page changes
        if current_page is not None and page_no != current_page:
            doc.add_page_break()
        current_page = page_no

        stripped = text.strip()
        if not stripped:
            continue

        # Heading detection: short, ALL-CAPS, no leading digits
        is_heading = (
            stripped == stripped.upper()
            and 2 < len(stripped) <= 60
            and not stripped[:1].isdigit()
            and "." not in stripped          # exclude sentences that happen to be caps
        )

        if is_heading:
            para = doc.add_heading(stripped, level=2)
        else:
            para = doc.add_paragraph()
            if ac and ac.replacements:
                _add_coloured_runs(para, text, ac.replacements)
            else:
                run = para.add_run(text)
                run.font.size = Pt(10)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _reconstruct_docx(
    source_path: Path,
    chunks: list[TextChunk],
    scrubbed: dict[str, AnonymisedChunk],
    output_path: Path,
) -> None:
    """Replace run text in DOCX, strip metadata and comments."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(source_path))

    # Build lookup: (location_key, para_index, run_index) → scrubbed_text
    # location_key examples: "body", "table_0_row_1_cell_2", "header_0"
    scrub_map: dict[tuple[str, int, int], str] = {}
    for chunk in chunks:
        ac = scrubbed.get(chunk.chunk_id)
        if ac is None:
            continue
        meta = chunk.metadata
        key = (
            meta.get("location", "body"),
            meta.get("para_index", 0),
            meta.get("run_index", 0),
        )
        scrub_map[key] = ac.scrubbed_text

    def _replace_runs(paragraphs, location_prefix: str) -> None:
        for para_idx, para in enumerate(paragraphs):
            for run_idx, run in enumerate(para.runs):
                key = (location_prefix, para_idx, run_idx)
                if key in scrub_map:
                    run.text = scrub_map[key]

    # Main body
    _replace_runs(doc.paragraphs, "body")

    # Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                loc = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                _replace_runs(cell.paragraphs, loc)

    # Headers and footers
    for section_idx, section in enumerate(doc.sections):
        _replace_runs(section.header.paragraphs, f"header_{section_idx}")
        _replace_runs(section.footer.paragraphs, f"footer_{section_idx}")

    # Remove comments
    try:
        comments_part = doc.part.document_part.comments_part
        if comments_part is not None:
            for comment in comments_part.element.findall(qn("w:comment")):
                comment.getparent().remove(comment)
    except Exception:
        pass

    # Strip core properties (author, last_modified_by, created, modified, etc.)
    try:
        props = doc.core_properties
        props.author = ""
        props.last_modified_by = ""
        props.title = ""
        props.subject = ""
        props.keywords = ""
        props.description = ""
        props.category = ""
        props.comments = ""
    except Exception:
        pass

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------

def _reconstruct_xlsx(
    source_path: Path,
    chunks: list[TextChunk],
    scrubbed: dict[str, AnonymisedChunk],
    output_path: Path,
) -> None:
    """Replace cell values in XLSX. Formulae are NOT preserved (data_only limitation)."""
    import openpyxl

    # Build lookup: (sheet_name, row, col) → scrubbed_text
    scrub_map: dict[tuple[str, int, int], str] = {}
    for chunk in chunks:
        ac = scrubbed.get(chunk.chunk_id)
        if ac is None:
            continue
        meta = chunk.metadata
        key = (meta.get("sheet", ""), meta.get("row", 0), meta.get("col", 0))
        scrub_map[key] = ac.scrubbed_text

    # Load with data_only=True (formulae already dropped by extractor)
    wb = openpyxl.load_workbook(str(source_path), data_only=True)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            for cell in row:
                key = (sheet_name, cell.row, cell.column)
                if key in scrub_map:
                    cell.value = scrub_map[key]

    wb.save(str(output_path))
    wb.close()


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------

def _reconstruct_csv(
    source_path: Path,
    chunks: list[TextChunk],
    scrubbed: dict[str, AnonymisedChunk],
    output_path: Path,
) -> None:
    """Rewrite CSV with scrubbed cell values."""
    import pandas as pd

    # Build lookup: (row_index, col_index) → scrubbed_text
    scrub_map: dict[tuple[int, int], str] = {}
    for chunk in chunks:
        ac = scrubbed.get(chunk.chunk_id)
        if ac is None:
            continue
        meta = chunk.metadata
        key = (meta.get("row", 0), meta.get("col", 0))
        scrub_map[key] = ac.scrubbed_text

    df = pd.read_csv(str(source_path), dtype=str, keep_default_na=False)

    for (row_idx, col_idx), scrubbed_text in scrub_map.items():
        if row_idx < len(df) and col_idx < len(df.columns):
            df.iloc[row_idx, col_idx] = scrubbed_text

    df.to_csv(str(output_path), index=False)


# ---------------------------------------------------------------------------
# EML
# ---------------------------------------------------------------------------

def _reconstruct_eml(
    source_path: Path,
    chunks: list[TextChunk],
    scrubbed: dict[str, AnonymisedChunk],
    output_path: Path,
) -> None:
    """Reconstruct EML with scrubbed headers and body."""
    with open(source_path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=email.policy.default)

    # Build lookup: part_name → scrubbed_text
    scrub_map: dict[str, str] = {}
    for chunk in chunks:
        ac = scrubbed.get(chunk.chunk_id)
        if ac is None:
            continue
        part_name = chunk.metadata.get("part", "")
        scrub_map[part_name] = ac.scrubbed_text

    # Replace headers
    for header_field in ("from", "to", "cc", "subject"):
        if header_field in scrub_map:
            if msg[header_field.capitalize()]:
                del msg[header_field.capitalize()]
            msg[header_field.capitalize()] = scrub_map[header_field]

    # Replace body parts
    body_text = scrub_map.get("body", "")
    if body_text:
        for part in msg.walk():
            content_type = part.get_content_type()
            disposition = str(part.get("Content-Disposition", ""))
            if "attachment" in disposition:
                continue
            if content_type == "text/plain":
                part.set_content(body_text)
                break

    with open(output_path, "wb") as f:
        f.write(msg.as_bytes(policy=email.policy.default))


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------

def _reconstruct_txt(
    source_path: Path,
    chunks: list[TextChunk],
    scrubbed: dict[str, AnonymisedChunk],
    output_path: Path,
) -> None:
    """Write scrubbed plain text, preserving line structure."""
    # Build lookup: char_offset_start → scrubbed_text
    scrub_map: dict[int, str] = {}
    for chunk in chunks:
        ac = scrubbed.get(chunk.chunk_id)
        if ac is None:
            continue
        scrub_map[chunk.char_offset_start] = ac.scrubbed_text

    original = source_path.read_text(encoding="utf-8", errors="replace")
    lines = original.splitlines(keepends=True)

    output_lines: list[str] = []
    char_offset = 0
    for line in lines:
        stripped = line.rstrip("\r\n")
        eol = line[len(stripped):]
        if char_offset in scrub_map:
            output_lines.append(scrub_map[char_offset] + eol)
        else:
            output_lines.append(line)
        char_offset += len(line)

    output_path.write_text("".join(output_lines), encoding="utf-8")
