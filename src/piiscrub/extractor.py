"""Document extraction layer.

Pulls raw text from each supported file format while tracking position
metadata needed by output.py to reconstruct the scrubbed document.

IMPORTANT: The actual filename is NEVER logged — only sha256(filename).

Supported formats: .pdf, .docx, .txt, .csv, .xlsx, .eml
Unsupported formats: raises UnsupportedFormatError

Coordinate systems by format:
  PDF  — TextChunk.metadata = {"page": int, "bbox": (x0, y0, x1, y1), "block_no": int}
  DOCX — TextChunk.metadata = {"para_index": int, "run_index": int, "location": str}
  XLSX — TextChunk.metadata = {"sheet": str, "row": int, "col": int}
  CSV  — TextChunk.metadata = {"row": int, "col": int, "col_name": str}
  EML  — TextChunk.metadata = {"part": "subject"|"from"|"to"|"cc"|"body"}
  TXT  — TextChunk.metadata = {"line_start": int}
"""

from __future__ import annotations

import email
import gc
import hashlib
from email import policy as email_policy
from pathlib import Path
from typing import Callable

from piiscrub._logging import get_logger
from piiscrub.models import TextChunk

_log = get_logger("extractor")

_SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx", ".eml"}


class UnsupportedFormatError(ValueError):
    """Raised when a file format is not supported by PIIScrub."""


def _hash_path(path: Path) -> str:
    return hashlib.sha256(str(path.resolve()).encode("utf-8")).hexdigest()


def extract(path: Path) -> list[TextChunk]:
    """Extract text chunks from a document file.

    Returns an ordered list of TextChunk objects. Each chunk maps to a
    discrete unit in the original document (page, paragraph, cell, etc.)

    Raises UnsupportedFormatError for unsupported file extensions.
    """
    path = Path(path)
    ext = path.suffix.lower()

    if ext not in _SUPPORTED_EXTENSIONS:
        raise UnsupportedFormatError(
            f"Unsupported file format: '{ext}'. "
            f"Supported formats: {', '.join(sorted(_SUPPORTED_EXTENSIONS))}"
        )

    _dispatch: dict[str, Callable[[Path], list[TextChunk]]] = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".csv": _extract_csv,
        ".eml": _extract_eml,
        ".txt": _extract_txt,
    }

    source_hash = _hash_path(path)
    chunks = _dispatch[ext](path)

    # Log stats without exposing the filename
    total_chars = sum(len(c.text) for c in chunks)
    _log.info(
        "extraction_complete",
        source_hash=source_hash[:12] + "...",  # abbreviated — not enough to reconstruct
        format=ext.lstrip("."),
        chunk_count=len(chunks),
        total_chars=total_chars,
    )

    gc.collect()
    return chunks


def _extract_pdf(path: Path) -> list[TextChunk]:
    """Extract text from PDF using PyMuPDF.

    Uses get_text("dict") to obtain per-span bounding boxes needed
    by output.py for spatial redaction. Also strips document metadata.
    """
    import fitz  # PyMuPDF

    source_hash = _hash_path(path)
    chunks: list[TextChunk] = []

    doc = fitz.open(str(path))
    # Strip metadata immediately on load (defensive — final strip happens in output.py)
    doc.set_metadata({})

    for page_no, page in enumerate(doc):
        raw = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        char_offset = 0

        for block in raw.get("blocks", []):
            if block.get("type") != 0:  # 0 = text block
                continue
            for line in block.get("lines", []):
                line_text = ""
                line_bbox = line.get("bbox", (0, 0, 0, 0))
                for span in line.get("spans", []):
                    line_text += span.get("text", "")

                if not line_text.strip():
                    char_offset += len(line_text)
                    continue

                chunk = TextChunk.create(
                    source_path_hash=source_hash,
                    fmt="pdf",
                    text=line_text,
                    page_or_sheet=page_no,
                    char_offset_start=char_offset,
                    metadata={
                        "page": page_no,
                        "bbox": line_bbox,
                        "block_no": block.get("number", 0),
                    },
                )
                chunks.append(chunk)
                char_offset += len(line_text)

    doc.close()
    return chunks


def _extract_docx(path: Path) -> list[TextChunk]:
    """Extract text from DOCX, including tables, headers, footers.

    Checks for tracked changes (w:ins/w:del) and warns if found.
    Deletes comments from the document part.
    """
    from docx import Document
    from docx.oxml.ns import qn

    source_hash = _hash_path(path)
    chunks: list[TextChunk] = []
    char_offset = 0

    doc = Document(str(path))

    # Warn if tracked changes are present (PII may be hidden in revisions)
    body_xml = doc.element.body.xml
    if "<w:ins " in body_xml or "<w:del " in body_xml:
        _log.warning(
            "docx_tracked_changes_detected",
            source_hash=source_hash[:12] + "...",
            warning="Document contains tracked changes. "
                    "Deleted text may contain PII that cannot be reliably scrubbed. "
                    "Accept/reject all changes before processing.",
        )

    # Remove comments
    try:
        comments_part = doc.part.document_part.comments_part
        if comments_part is not None:
            for comment in comments_part.element.findall(qn("w:comment")):
                comment.getparent().remove(comment)
    except Exception:
        pass  # No comments part present

    def _add_paragraph_chunk(para_text: str, para_index: int, run_index: int, location: str) -> None:
        nonlocal char_offset
        if not para_text.strip():
            char_offset += len(para_text)
            return
        chunk = TextChunk.create(
            source_path_hash=source_hash,
            fmt="docx",
            text=para_text,
            page_or_sheet=None,
            char_offset_start=char_offset,
            metadata={
                "para_index": para_index,
                "run_index": run_index,
                "location": location,
            },
        )
        chunks.append(chunk)
        char_offset += len(para_text)

    # Main body paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            if run.text:
                _add_paragraph_chunk(run.text, para_idx, run_idx, "body")

    # Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(para.runs):
                        if run.text:
                            _add_paragraph_chunk(
                                run.text,
                                para_idx,
                                run_idx,
                                f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}",
                            )

    # Headers and footers
    for section_idx, section in enumerate(doc.sections):
        for header_para in section.header.paragraphs:
            for run_idx, run in enumerate(header_para.runs):
                if run.text:
                    _add_paragraph_chunk(run.text, 0, run_idx, f"header_{section_idx}")
        for footer_para in section.footer.paragraphs:
            for run_idx, run in enumerate(footer_para.runs):
                if run.text:
                    _add_paragraph_chunk(run.text, 0, run_idx, f"footer_{section_idx}")

    return chunks


def _extract_xlsx(path: Path) -> list[TextChunk]:
    """Extract text from XLSX using openpyxl with data_only=True.

    NOTE: data_only=True means formula expressions are not extracted,
    only cached display values. Formulas will NOT be preserved in output.
    This is documented as a known limitation.
    """
    import openpyxl

    source_hash = _hash_path(path)
    chunks: list[TextChunk] = []
    char_offset = 0

    wb = openpyxl.load_workbook(str(path), data_only=True)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                cell_text = str(cell.value)
                if not cell_text.strip():
                    continue
                chunk = TextChunk.create(
                    source_path_hash=source_hash,
                    fmt="xlsx",
                    text=cell_text,
                    page_or_sheet=wb.sheetnames.index(sheet_name),
                    char_offset_start=char_offset,
                    metadata={
                        "sheet": sheet_name,
                        "row": cell.row,
                        "col": cell.column,
                        "coordinate": cell.coordinate,
                    },
                )
                chunks.append(chunk)
                char_offset += len(cell_text)

    wb.close()
    return chunks


def _extract_csv(path: Path) -> list[TextChunk]:
    """Extract text from CSV row by row using pandas."""
    import pandas as pd

    source_hash = _hash_path(path)
    chunks: list[TextChunk] = []
    char_offset = 0

    df = pd.read_csv(str(path), dtype=str, keep_default_na=False)

    for row_idx, row in df.iterrows():
        for col_name in df.columns:
            cell_text = str(row[col_name]).strip()
            if not cell_text:
                continue
            chunk = TextChunk.create(
                source_path_hash=source_hash,
                fmt="csv",
                text=cell_text,
                page_or_sheet=None,
                char_offset_start=char_offset,
                metadata={
                    "row": int(row_idx),  # type: ignore[arg-type]
                    "col": df.columns.tolist().index(col_name),
                    "col_name": col_name,
                },
            )
            chunks.append(chunk)
            char_offset += len(cell_text)

    return chunks


def _extract_eml(path: Path) -> list[TextChunk]:
    """Extract text from EML files.

    Extracts: From, To, CC, Subject headers + all text/* MIME body parts.
    Binary attachments are skipped (v0.1 limitation — logged as a warning).
    """
    source_hash = _hash_path(path)
    chunks: list[TextChunk] = []
    char_offset = 0

    with open(path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=email_policy.default)

    def _add_header_chunk(value: str, part_name: str) -> None:
        nonlocal char_offset
        value = str(value).strip()
        if not value:
            return
        chunk = TextChunk.create(
            source_path_hash=source_hash,
            fmt="eml",
            text=value,
            page_or_sheet=None,
            char_offset_start=char_offset,
            metadata={"part": part_name},
        )
        chunks.append(chunk)
        char_offset += len(value)

    _add_header_chunk(msg.get("From", ""), "from")
    _add_header_chunk(msg.get("To", ""), "to")
    _add_header_chunk(msg.get("CC", ""), "cc")
    _add_header_chunk(msg.get("Subject", ""), "subject")

    attachment_count = 0
    for part in msg.walk():
        content_type = part.get_content_type()
        disposition = str(part.get("Content-Disposition", ""))

        if "attachment" in disposition:
            attachment_count += 1
            continue

        if content_type in ("text/plain", "text/html"):
            try:
                body = part.get_content()
                if body and body.strip():
                    chunk = TextChunk.create(
                        source_path_hash=source_hash,
                        fmt="eml",
                        text=body,
                        page_or_sheet=None,
                        char_offset_start=char_offset,
                        metadata={"part": "body", "content_type": content_type},
                    )
                    chunks.append(chunk)
                    char_offset += len(body)
            except Exception:
                pass

    if attachment_count > 0:
        _log.warning(
            "eml_attachments_skipped",
            source_hash=source_hash[:12] + "...",
            attachment_count=attachment_count,
            note="Binary attachment PII detection is out of scope for v0.1",
        )

    return chunks


def _extract_txt(path: Path) -> list[TextChunk]:
    """Extract text from plain text files, one chunk per line."""
    source_hash = _hash_path(path)
    chunks: list[TextChunk] = []
    char_offset = 0

    text = path.read_text(encoding="utf-8", errors="replace")

    # Split into lines but keep line endings for accurate offsets
    lines = text.splitlines(keepends=True)
    for line in lines:
        stripped = line.rstrip("\r\n")
        if stripped.strip():
            chunk = TextChunk.create(
                source_path_hash=source_hash,
                fmt="txt",
                text=stripped,
                page_or_sheet=None,
                char_offset_start=char_offset,
                metadata={"line_start": char_offset},
            )
            chunks.append(chunk)
        char_offset += len(line)

    return chunks
